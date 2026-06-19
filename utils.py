"""utils.py — Document processing and helper utilities."""

import io, re, json, logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        parts = []
        for i, page in enumerate(reader.pages):
            t = page.extract_text()
            if t:
                parts.append(f"--- Page {i+1} ---\n{t}")
        return "\n\n".join(parts) if parts else "No text extracted from PDF."
    except Exception as e:
        return f"PDF extraction error: {e}"


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts) if parts else "No text extracted from DOCX."
    except Exception as e:
        return f"DOCX extraction error: {e}"


def extract_text_from_txt(file_bytes: bytes) -> str:
    for enc in ["utf-8", "latin-1", "cp1252"]:
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def process_uploaded_file(uploaded_file) -> tuple[str, str]:
    data = uploaded_file.read()
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(data), "PDF"
    elif name.endswith((".docx", ".doc")):
        return extract_text_from_docx(data), "DOCX"
    elif name.endswith(".txt"):
        return extract_text_from_txt(data), "TXT"
    elif name.endswith(".md"):
        return extract_text_from_txt(data), "Markdown"
    else:
        try:
            return extract_text_from_txt(data), "Text"
        except Exception:
            return "Unable to extract text from this file type.", "Unknown"


def truncate_text(text: str, max_chars: int = 12000) -> str:
    if len(text) <= max_chars:
        return text
    truncated = text[:max_chars]
    last_period = truncated.rfind(".")
    if last_period > max_chars * 0.8:
        truncated = truncated[:last_period + 1]
    return truncated + f"\n\n[TRUNCATED — showing first {max_chars:,} of {len(text):,} chars]"


def parse_json_response(response_text: str) -> Optional[dict]:
    text = response_text.strip()
    # Strip markdown fences
    if "```json" in text:
        m = re.search(r"```json\s*(.*?)\s*```", text, re.DOTALL)
        if m:
            text = m.group(1)
    elif "```" in text:
        m = re.search(r"```\s*(.*?)\s*```", text, re.DOTALL)
        if m:
            text = m.group(1)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    try:
        s, e = text.find("{"), text.rfind("}")
        if s != -1 and e != -1:
            return json.loads(text[s:e+1])
    except json.JSONDecodeError:
        pass
    logger.error(f"JSON parse failed. First 300 chars: {text[:300]}")
    return None


def validate_analysis_structure(analysis: dict) -> tuple[bool, list]:
    required = ["project_type", "identified_hazards", "risk_levels",
                "safety_checklist", "compliance_requirements",
                "recommended_controls", "overall_risk_assessment", "summary"]
    missing = [f for f in required if f not in analysis]
    return len(missing) == 0, missing


def get_risk_emoji(risk_level: str) -> str:
    return {"critical": "🔴", "high": "🟠", "medium": "🟡", "low": "🟢"}.get(
        risk_level.lower(), "⚪"
    )


def get_domain_icon(domain: str) -> str:
    icons = {
        "Mechanical Engineering": "⚙️", "Electrical Engineering": "⚡",
        "Electronics Engineering": "🔌", "Civil Engineering": "🏗️",
        "Software Engineering": "💻", "Industrial Automation": "🤖",
        "Manufacturing Systems": "🏭", "Robotics": "🦾",
        "IoT Systems": "📡", "Construction Projects": "🏚️",
    }
    return icons.get(domain, "🔧")
